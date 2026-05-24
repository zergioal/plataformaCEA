# -*- coding: utf-8 -*-
# Generador de Guia Didactica - CEA Madre Maria Oliva

import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

LQ = '“'
RQ = '”'

def q(text):
    return LQ + text + RQ

BASE_DIR  = os.path.dirname(os.path.abspath(__file__))
LOGO_PATH = os.path.join(BASE_DIR, 'src', 'assets', 'logo-cea.png')
FOTO_PATH = os.path.join(BASE_DIR, 'src', 'assets', 'CEA.jpeg')
OUT_PATH  = os.path.join(BASE_DIR, 'Guia_Didactica_CEA_Madre_Maria_Oliva.docx')

NAVY      = RGBColor(0x0d, 0x2b, 0x55)
BLUE      = RGBColor(0x1a, 0x6b, 0xb5)
TEAL      = RGBColor(0x0a, 0x7a, 0x6e)
GREEN_D   = RGBColor(0x14, 0x53, 0x2d)
AMBER     = RGBColor(0x92, 0x40, 0x0e)
GRAY_TEXT = RGBColor(0x37, 0x41, 0x51)
WHITE     = RGBColor(0xFF, 0xFF, 0xFF)

BG_ACTIVITY = 'E8F5E9'
BG_TIP      = 'FFF8DC'
BG_SCREEN   = 'EEF2F7'
BG_NAVY     = '0D2B55'
BG_BLUE     = '1A6BB5'
BG_TEAL     = '0A7A6E'

FONT_BODY  = 'Calibri'
FONT_TITLE = 'Calibri'


def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)


def set_table_border(table, hex_color='CCCCCC', sz=4):
    tbl   = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'),   'single')
        el.set(qn('w:sz'),    str(sz))
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), hex_color)
        tblBorders.append(el)
    tblPr.append(tblBorders)


def set_para_shading(para, hex_color):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    pPr.append(shd)


def set_col_width(table, col_idx, width_cm):
    for row in table.rows:
        row.cells[col_idx].width = Cm(width_cm)


def page_break(doc):
    p   = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    run = p.add_run()
    br  = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    run._r.append(br)


def heading_para(doc, text, level, color=None, bg=None, size=None, bold=True,
                 space_before=16, space_after=6, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if align:
        p.alignment = align
    if bg:
        set_para_shading(p, bg)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(4)
        pPr = p._p.get_or_add_pPr()
        ind = OxmlElement('w:ind')
        ind.set(qn('w:left'), '160')
        pPr.append(ind)
    run = p.add_run(text)
    run.bold = bold
    run.font.name  = FONT_TITLE
    run.font.size  = Pt(size or (20 if level == 0 else 15 if level == 1 else 13 if level == 2 else 11))
    run.font.color.rgb = color or NAVY
    return p


def body_para(doc, text='', bold=False, italic=False, size=11, color=None,
              indent_cm=0, space_after=4, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    if indent_cm:
        p.paragraph_format.left_indent = Cm(indent_cm)
    if align:
        p.alignment = align
    if text:
        run = p.add_run(text)
        run.bold       = bold
        run.italic     = italic
        run.font.name  = FONT_BODY
        run.font.size  = Pt(size)
        if color:
            run.font.color.rgb = color
    return p


def numbered_step(doc, number, text, indent_cm=0.5):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.left_indent  = Cm(indent_cm)
    rn = p.add_run(f'{number}.  ')
    rn.bold           = True
    rn.font.name      = FONT_BODY
    rn.font.size      = Pt(11)
    rn.font.color.rgb = BLUE
    rt = p.add_run(text)
    rt.font.name = FONT_BODY
    rt.font.size = Pt(11)
    return p


def bullet_para(doc, text, indent_cm=0.7, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.left_indent  = Cm(indent_cm)
    rb = p.add_run('• ')
    rb.font.name      = FONT_BODY
    rb.font.size      = Pt(11)
    rb.font.color.rgb = BLUE
    if bold_prefix:
        rp = p.add_run(bold_prefix)
        rp.bold      = True
        rp.font.name = FONT_BODY
        rp.font.size = Pt(11)
    rt = p.add_run(text)
    rt.font.name = FONT_BODY
    rt.font.size = Pt(11)
    return p


def tip_box(doc, text, icon='\U0001f4a1', title='Consejo'):
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_border(t, hex_color='F5C842', sz=6)
    c = t.cell(0, 0)
    set_cell_bg(c, BG_TIP)
    p = c.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pPr = p._p.get_or_add_pPr()
    ind = OxmlElement('w:ind')
    ind.set(qn('w:left'), '120')
    pPr.append(ind)
    r1 = p.add_run(f'{icon} {title}: ')
    r1.bold           = True
    r1.font.name      = FONT_BODY
    r1.font.size      = Pt(10)
    r1.font.color.rgb = AMBER
    r2 = p.add_run(text)
    r2.font.name      = FONT_BODY
    r2.font.size      = Pt(10)
    r2.font.color.rgb = GRAY_TEXT
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def activity_box(doc, title, steps):
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_border(t, hex_color='2E7D32', sz=8)
    c = t.cell(0, 0)
    set_cell_bg(c, BG_ACTIVITY)
    p0 = c.paragraphs[0]
    p0.paragraph_format.space_before = Pt(4)
    p0.paragraph_format.space_after  = Pt(4)
    pPr = p0._p.get_or_add_pPr()
    ind = OxmlElement('w:ind')
    ind.set(qn('w:left'), '120')
    pPr.append(ind)
    r = p0.add_run('✏️  ' + title)
    r.bold           = True
    r.font.name      = FONT_BODY
    r.font.size      = Pt(11)
    r.font.color.rgb = GREEN_D
    for i, step in enumerate(steps, 1):
        pc = c.add_paragraph()
        pc.paragraph_format.space_before = Pt(2)
        pc.paragraph_format.space_after  = Pt(2)
        pPr2 = pc._p.get_or_add_pPr()
        ind2 = OxmlElement('w:ind')
        ind2.set(qn('w:left'), '240')
        pPr2.append(ind2)
        rn = pc.add_run(f'{i}.  ')
        rn.bold           = True
        rn.font.name      = FONT_BODY
        rn.font.size      = Pt(10)
        rn.font.color.rgb = GREEN_D
        rt = pc.add_run(step)
        rt.font.name      = FONT_BODY
        rt.font.size      = Pt(10)
        rt.font.color.rgb = GRAY_TEXT
    pf = c.add_paragraph()
    pf.paragraph_format.space_after = Pt(2)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t


def screenshot_box(doc, label, height_cm=4.5):
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_border(t, hex_color='9EAFC2', sz=6)
    c = t.cell(0, 0)
    set_cell_bg(c, BG_SCREEN)
    c.height = Cm(height_cm)
    c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = c.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p.add_run('\U0001f4f7  ')
    r1.font.size = Pt(18)
    r2 = p.add_run(label)
    r2.italic         = True
    r2.font.name      = FONT_BODY
    r2.font.size      = Pt(9)
    r2.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t


def section_divider(doc, text, bg_hex, text_color=None):
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    c = t.cell(0, 0)
    set_cell_bg(c, bg_hex)
    p = c.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    r = p.add_run(text)
    r.bold           = True
    r.font.name      = FONT_TITLE
    r.font.size      = Pt(14)
    r.font.color.rgb = text_color or WHITE
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(6)
    return t


def styled_table(doc, headers, rows, col_widths=None,
                 header_bg=None, alt_row_bg='F0F4FA', font_size=10):
    if header_bg is None:
        header_bg = BG_NAVY
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_border(table, hex_color='B0BEC5', sz=4)
    for j, h in enumerate(headers):
        c = table.cell(0, j)
        set_cell_bg(c, header_bg)
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(str(h))
        r.bold           = True
        r.font.name      = FONT_BODY
        r.font.size      = Pt(font_size)
        r.font.color.rgb = WHITE
    for i, row_data in enumerate(rows):
        bg = alt_row_bg if i % 2 == 0 else 'FFFFFF'
        for j, cell_text in enumerate(row_data):
            c = table.cell(i + 1, j)
            set_cell_bg(c, bg)
            p = c.paragraphs[0]
            text = str(cell_text)
            if text in ('✅', '❌', '—'):
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(text)
            r.font.name = FONT_BODY
            r.font.size = Pt(font_size)
    if col_widths:
        for col_i, w in enumerate(col_widths):
            set_col_width(table, col_i, w)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table


# ══════════════════════════════════════════════════════════════════════════════
#  CONSTRUCCION DEL DOCUMENTO
# ══════════════════════════════════════════════════════════════════════════════

def build_document():
    doc = Document()

    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    style = doc.styles['Normal']
    style.font.name  = FONT_BODY
    style.font.size  = Pt(11)
    style.paragraph_format.space_after  = Pt(6)
    style.paragraph_format.space_before = Pt(0)

    # ── PORTADA ───────────────────────────────────────────────────────────────
    if os.path.exists(FOTO_PATH):
        try:
            p_img = doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_img.paragraph_format.space_before = Pt(0)
            p_img.paragraph_format.space_after  = Pt(0)
            p_img.add_run().add_picture(FOTO_PATH, width=Inches(6.2))
        except Exception:
            pass

    cover_table = doc.add_table(rows=1, cols=2)
    cover_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_border(cover_table, hex_color='0D2B55', sz=0)
    cl = cover_table.cell(0, 0)
    cr = cover_table.cell(0, 1)
    set_cell_bg(cl, BG_NAVY)
    set_cell_bg(cr, BG_NAVY)
    cl.width = Cm(3.0)
    cr.width = Cm(13.5)
    cl.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    cr.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    if os.path.exists(LOGO_PATH):
        try:
            pl = cl.paragraphs[0]
            pl.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pl.paragraph_format.space_before = Pt(8)
            pl.paragraph_format.space_after  = Pt(8)
            pl.add_run().add_picture(LOGO_PATH, width=Cm(2.4))
        except Exception:
            pass

    pr = cr.paragraphs[0]
    pr.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pr.paragraph_format.space_before = Pt(10)
    pr.paragraph_format.space_after  = Pt(2)
    pPr = pr._p.get_or_add_pPr()
    ind = OxmlElement('w:ind')
    ind.set(qn('w:left'), '160')
    pPr.append(ind)
    r1 = pr.add_run('CEA “Madre Maria Oliva”')
    r1.bold           = True
    r1.font.name      = FONT_TITLE
    r1.font.size      = Pt(13)
    r1.font.color.rgb = WHITE

    pr2 = cr.add_paragraph()
    pr2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pPr2 = pr2._p.get_or_add_pPr()
    ind2 = OxmlElement('w:ind')
    ind2.set(qn('w:left'), '160')
    pPr2.append(ind2)
    r2 = pr2.add_run('Centro de Educación Alternativa')
    r2.font.name      = FONT_BODY
    r2.font.size      = Pt(10)
    r2.font.color.rgb = RGBColor(0xAA, 0xC4, 0xE0)
    pr2.paragraph_format.space_after = Pt(10)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    pt = doc.add_paragraph()
    pt.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pt.paragraph_format.space_before = Pt(16)
    pt.paragraph_format.space_after  = Pt(4)
    rt = pt.add_run('GUÍA PRÁCTICA')
    rt.bold           = True
    rt.font.name      = FONT_TITLE
    rt.font.size      = Pt(32)
    rt.font.color.rgb = NAVY

    ps = doc.add_paragraph()
    ps.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ps.paragraph_format.space_after = Pt(4)
    rs = ps.add_run('Plataforma Digital Institucional')
    rs.bold           = True
    rs.font.name      = FONT_TITLE
    rs.font.size      = Pt(20)
    rs.font.color.rgb = BLUE

    ld = doc.add_paragraph()
    ld.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ld.paragraph_format.space_after = Pt(16)
    rld = ld.add_run('─' * 36)
    rld.font.color.rgb = BLUE
    rld.font.size = Pt(12)

    pst = doc.add_paragraph()
    pst.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pst.paragraph_format.space_after = Pt(6)
    rst = pst.add_run('Para uso de Docentes, Secretaria y Directora')
    rst.font.name      = FONT_BODY
    rst.font.size      = Pt(13)
    rst.font.color.rgb = GRAY_TEXT

    pv = doc.add_paragraph()
    pv.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pv.paragraph_format.space_after = Pt(4)
    rv = pv.add_run('Versión 1.0  ·  Mayo 2026  ·  ceamadremariaoliva.edu.bo')
    rv.font.name      = FONT_BODY
    rv.font.size      = Pt(10)
    rv.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    page_break(doc)

    # ── PRESENTACION ──────────────────────────────────────────────────────────
    heading_para(doc, 'Presentación', level=1, color=NAVY, size=16, space_before=4)
    body_para(doc,
        'Esta guía ha sido elaborada para que el personal docente y administrativo del '
        'CEA Madre María Oliva pueda utilizar la plataforma digital de forma eficiente '
        'en su trabajo cotidiano.')
    body_para(doc,
        'La plataforma permite gestionar estudiantes, registrar asistencia, ingresar '
        'calificaciones, publicar material de estudio y administrar la institución '
        'desde cualquier dispositivo con acceso a internet.')
    body_para(doc,
        'No se requiere experiencia previa. Cada sección incluye los pasos necesarios '
        'y una actividad práctica para afianzar lo aprendido.')

    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    styled_table(doc,
        headers=['Módulo', 'Dirigido a', 'Contenido principal'],
        rows=[
            ['Módulo 0', 'Todos',      'Primeros pasos: acceso, roles, contraseña'],
            ['Módulo 1', 'Docentes',   'Estudiantes, asistencia, notas, contenido, IA'],
            ['Módulo 2', 'Secretaria', 'Consulta y exportación de reportes PDF'],
            ['Módulo 3', 'Directora',  'Gestión completa de la institución'],
        ],
        col_widths=[3.0, 3.5, 9.0],
        font_size=10,
    )

    page_break(doc)

    # ══════════════════════════════════════════════════════════════════════════
    #  MODULO 0
    # ══════════════════════════════════════════════════════════════════════════
    section_divider(doc, 'MÓDULO 0 — Primeros Pasos  ·  Para todos los usuarios', BG_NAVY)

    heading_para(doc, '0.1  ¿Qué es la plataforma?', level=2, color=BLUE, size=13, space_before=10)
    body_para(doc,
        'La plataforma digital del CEA Madre María Oliva es un sistema web institucional '
        'accesible desde: https://ceamadremariaoliva.edu.bo/')
    body_para(doc, 'Tiene dos partes:', bold=True)
    bullet_para(doc,
        'Información institucional, carreras, galería, requisitos de inscripción y contacto.',
        bold_prefix='Parte pública — ')
    bullet_para(doc,
        'Panel privado para docentes, secretaria, directora y estudiantes.',
        bold_prefix='Parte privada — ')

    screenshot_box(doc, 'Captura: Página de inicio pública del CEA', height_cm=4.0)

    heading_para(doc, '0.2  Cómo ingresar al sistema', level=2, color=BLUE, size=13)
    for i, step in enumerate([
        'Abrir el navegador (Chrome o Firefox recomendados).',
        'Ir a: https://ceamadremariaoliva.edu.bo/',
        'Hacer clic en el botón “Iniciar sesión”.',
        'Ingresar tu código de usuario (ej. DOC-001) y tu contraseña.',
        'Hacer clic en “Ingresar”.',
    ], 1):
        numbered_step(doc, i, step)

    screenshot_box(doc, 'Captura: Pantalla de inicio de sesión', height_cm=3.5)
    tip_box(doc,
        'Si ingresas una contraseña incorrecta varias veces, tu cuenta se bloqueará automáticamente. '
        'Comunícate con la directora para que la desbloquee.',
        icon='⚠️', title='Importante')

    heading_para(doc, '0.3  Roles del sistema', level=2, color=BLUE, size=13)
    styled_table(doc,
        headers=['Rol', 'Descripción de acceso'],
        rows=[
            ['Docente',     'Panel de trabajo: sus estudiantes, asistencia, notas y contenido'],
            ['Secretaria',  'Panel administrativo en modo consulta (no crea ni elimina)'],
            ['Directora',   'Acceso completo a todas las funciones del sistema'],
            ['Estudiante',  'Sus módulos, calificaciones y registro de asistencia personal'],
        ],
        col_widths=[3.5, 12.0],
        font_size=10,
    )

    heading_para(doc, '0.4  Cambiar la contraseña temporal', level=2, color=BLUE, size=13)
    body_para(doc,
        'Al recibir tu cuenta tendrás una contraseña temporal. '
        'Es importante cambiarla en el primer ingreso.')
    for i, step in enumerate([
        'Ingresar al sistema con la contraseña temporal.',
        'En tu panel, buscar la sección “Cambiar contraseña” o “Seguridad”.',
        'Ingresar la contraseña actual (la temporal).',
        'Escribir la nueva contraseña — mínimo 6 caracteres.',
        'Confirmar la nueva contraseña y guardar.',
    ], 1):
        numbered_step(doc, i, step)
    tip_box(doc, 'Usa una contraseña que recuerdes pero que no sea obvia. Puedes combinar letras y números.')

    activity_box(doc,
        'Actividad 0 — Primer ingreso',
        [
            'Cada participante ingresa al sistema con sus credenciales.',
            'Cambiar la contraseña temporal por una personal.',
            'Verificar que el nombre aparece correctamente en el panel.',
        ]
    )

    page_break(doc)

    # ══════════════════════════════════════════════════════════════════════════
    #  MODULO 1
    # ══════════════════════════════════════════════════════════════════════════
    section_divider(doc, 'MÓDULO 1 — Guía para Docentes', BG_TEAL)

    # 1.1
    heading_para(doc, '1.1  Mi Perfil', level=1, color=TEAL, size=14, space_before=10)
    body_para(doc, 'Tu perfil contiene los datos personales visibles en el sistema. Mantenlos actualizados.')

    heading_para(doc, 'Editar mis datos', level=2, color=BLUE, size=12, space_before=8)
    for i, step in enumerate([
        'En tu panel principal, hacer clic en tu nombre o avatar.',
        'Actualizar teléfono/celular y correo electrónico de contacto.',
        'Seleccionar el grado académico (T.S., Lic., Ing., M.Sc., Dr.).',
        'Hacer clic en “Guardar”.',
    ], 1):
        numbered_step(doc, i, step)
    tip_box(doc,
        'El grado académico aparecerá como prefijo en tu nombre en los reportes PDF. '
        'Ej.: Lic. García Mamani, Juan.')

    heading_para(doc, 'Cambiar mi avatar', level=2, color=BLUE, size=12, space_before=8)
    body_para(doc,
        'El avatar es el ícono que te representa en el sistema. Para cambiarlo: '
        'hacer clic en el avatar actual → seleccionar un nuevo ícono → confirmar.')

    heading_para(doc, 'Cambiar mi contraseña', level=2, color=BLUE, size=12, space_before=8)
    for i, step in enumerate([
        'En tu perfil, buscar la opción “Cambiar contraseña”.',
        'Ingresar la contraseña actual.',
        'Escribir y confirmar la nueva contraseña (mínimo 6 caracteres).',
        'Guardar.',
    ], 1):
        numbered_step(doc, i, step)

    screenshot_box(doc, 'Captura: Sección de perfil del docente — datos personales y grado académico', height_cm=4.0)

    # 1.2
    heading_para(doc, '1.2  Mis Estudiantes', level=1, color=TEAL, size=14, space_before=14)
    body_para(doc, 'Muestra la lista de estudiantes de tu carrera y turno con información actualizada.')

    heading_para(doc, 'Ver la lista', level=2, color=BLUE, size=12, space_before=8)
    body_para(doc,
        'La tabla incluye: código, nombre completo, nivel, teléfono y última vez activo en la plataforma. '
        'Usa la barra de búsqueda para filtrar por nombre o código, y haz clic en los encabezados para ordenar.')

    heading_para(doc, 'Registrar un nuevo estudiante', level=2, color=BLUE, size=12, space_before=8)
    body_para(doc, 'Haz clic en “+ Añadir Estudiante” y completa el formulario:', size=11)

    body_para(doc, 'Datos obligatorios:', bold=True, size=10, color=GREEN_D, indent_cm=0.5)
    for campo in [
        'Nombre(s), Apellido paterno y/o materno (al menos uno).',
        'N° de Carnet de Identidad.',
        'Género (Masculino / Femenino).',
        'Fecha de nacimiento — el sistema valida la edad mínima de 14 años.',
        'Nivel al que pertenece.',
        'Contraseña temporal (se genera automáticamente, puedes modificarla).',
    ]:
        bullet_para(doc, campo, indent_cm=1.0)

    tip_box(doc,
        'El sistema genera automáticamente el código del estudiante (ej. SIS-042). '
        'Anota el código y la contraseña para entregárselos al estudiante.')

    heading_para(doc, 'Mesa Directiva', level=2, color=BLUE, size=12, space_before=8)
    body_para(doc,
        'Los miembros de la Mesa Directiva pueden ver la asistencia de todos los niveles de la carrera. '
        'Para asignar: abrir el perfil del estudiante → activar “Mesa Directiva” → guardar.')

    screenshot_box(doc, "Captura: Lista de estudiantes del docente con botón '+ Añadir Estudiante'", height_cm=4.0)
    screenshot_box(doc, 'Captura: Formulario de registro de nuevo estudiante', height_cm=4.5)

    activity_box(doc,
        'Actividad 1.2 — Registrar estudiantes',
        [
            'Cada docente registra en el sistema al menos 3 estudiantes reales de su grupo.',
            'Completar todos los campos obligatorios (carnet, género, fecha de nacimiento, nivel).',
            'Anotar el código y contraseña generados para cada estudiante.',
            'Verificar que aparecen en la lista y practicar la búsqueda por nombre.',
        ]
    )

    page_break(doc)

    # 1.3
    heading_para(doc, '1.3  Registro de Asistencia', level=1, color=TEAL, size=14, space_before=6)
    body_para(doc,
        'La asistencia alimenta automáticamente las dimensiones SER y DECIDIR en las calificaciones. '
        'Es importante mantenerla actualizada.')

    heading_para(doc, 'Estados de asistencia', level=2, color=BLUE, size=12, space_before=8)
    styled_table(doc,
        headers=['Código', 'Significado', 'Color en pantalla'],
        rows=[
            ['P',       'Presente',              'Verde'],
            ['A',       'Atraso',                'Amarillo'],
            ['F',       'Falta',                 'Rojo'],
            ['L',       'Licencia',              'Azul'],
            ['(vacío)', 'Sin registro aún', '—'],
        ],
        col_widths=[2.0, 4.0, 5.0],
        font_size=10,
    )

    heading_para(doc, 'Cómo registrar asistencia', level=2, color=BLUE, size=12, space_before=8)
    for i, step in enumerate([
        'Ir a la sección “Asistencia” en tu panel.',
        'Seleccionar el mes y el nivel del grupo.',
        'Se mostrará una tabla con los estudiantes en filas y los días hábiles (lunes a viernes) en columnas.',
        'Hacer clic en la celda del día y estudiante para cambiar el estado.',
        'El estado cambia en ciclo: vacío → P → A → F → L → vacío.',
        'Los cambios se guardan automáticamente.',
    ], 1):
        numbered_step(doc, i, step)

    body_para(doc,
        'Al final de cada fila verás el porcentaje de asistencia del mes: '
        'verde (≥80%), amarillo (60–79%), rojo (<60%).')

    screenshot_box(doc, 'Captura: Tabla de registro de asistencia por mes y nivel', height_cm=5.0)

    activity_box(doc,
        'Actividad 1.3 — Registrar asistencia',
        [
            'Registrar la asistencia de la semana actual para tu grupo.',
            'Verificar que los porcentajes de asistencia se calculan correctamente.',
            'Exportar el PDF de asistencia del mes actual (botón “Exportar PDF”).',
        ]
    )

    page_break(doc)

    # 1.4
    heading_para(doc, '1.4  Calificaciones — Vista del Módulo', level=1, color=TEAL, size=14, space_before=6)
    body_para(doc,
        'El sistema sigue el Modelo Educativo Sociocomunitario Productivo con cinco dimensiones evaluativas.')

    styled_table(doc,
        headers=['Dimensión', 'Rango', 'Descripción'],
        rows=[
            ['SER',                    '2 — 10 pts',   'Valores y convivencia. Calculado desde asistencia'],
            ['SABER',                  '6 — 30 pts',   'Conocimiento teórico y conceptual'],
            ['HACER Proceso',          '6 — 30 pts',   'Proceso práctico. Influenciado por avance en módulo'],
            ['HACER Producto',         '6 — 30 pts',   'Resultado y producto final'],
            ['DECIDIR',                '2 — 10 pts',   'Compromiso y participación. Calculado desde asistencia'],
            ['Autoevaluación SER',     '—', 'El estudiante evalúa sus propios valores'],
            ['Autoevaluación DECIDIR', '—', 'El estudiante evalúa su propio compromiso'],
            ['TOTAL',                  '60 — 100 pts', 'Suma de todas las dimensiones'],
        ],
        col_widths=[4.5, 3.0, 8.0],
        font_size=10,
    )

    tip_box(doc,
        'Los campos de nota no aceptan valores fuera del rango. '
        'Si ingresas un número incorrecto, el campo se sacudirá indicando el error.',
        icon='⚠️', title='Importante')

    heading_para(doc, 'Ingresar calificaciones', level=2, color=BLUE, size=12, space_before=8)
    for i, step in enumerate([
        'Ir a la sección “Calificaciones” en tu panel.',
        'Seleccionar el módulo a calificar.',
        'Para cada estudiante, ingresar los valores de cada dimensión.',
        'SER y DECIDIR se sugieren automáticamente según la asistencia.',
        'HACER Proceso muestra una sugerencia según el avance del estudiante en el módulo.',
        'El total se calcula automáticamente.',
        'Agregar una observación si es necesario.',
        'Hacer clic en “Guardar”.',
    ], 1):
        numbered_step(doc, i, step)

    screenshot_box(doc, 'Captura: Tabla de calificaciones del módulo con todas las dimensiones', height_cm=5.0)

    activity_box(doc,
        'Actividad 1.4 — Ingresar calificaciones',
        [
            'Seleccionar el módulo vigente de tu carrera.',
            'Ingresar las notas de al menos 3 estudiantes en todas las dimensiones.',
            'Verificar que el total se calcula automáticamente.',
            'Exportar las notas del módulo a PDF.',
        ]
    )

    page_break(doc)

    # 1.5
    heading_para(doc, '1.5  Calificaciones — Registro por Dimensión y Actividades',
                 level=1, color=TEAL, size=14, space_before=6)
    body_para(doc,
        'Además del ingreso general del módulo, puedes registrar las notas actividad por actividad '
        'dentro de cada dimensión, lo que da mayor detalle y transparencia a la evaluación.')

    heading_para(doc, '¿Qué son las actividades?', level=2, color=BLUE, size=12, space_before=8)
    body_para(doc,
        'Cada dimensión (SABER, HACER Proceso, HACER Producto, Autoevaluación SER, '
        'Autoevaluación DECIDIR) puede tener varias actividades asociadas, correspondientes '
        'a las secciones de contenido que creaste en el módulo.')

    heading_para(doc, 'Ingresar notas por actividad', level=2, color=BLUE, size=12, space_before=8)
    for i, step in enumerate([
        'Desde la vista de calificaciones del módulo, hacer clic en la dimensión específica (ej. “SABER”).',
        'Se abrirá la tabla de actividades con estudiantes en filas y actividades en columnas.',
        'Ingresar la nota de cada estudiante en cada actividad.',
        'El sistema suma automáticamente para obtener el total de la dimensión.',
        'Guardar.',
    ], 1):
        numbered_step(doc, i, step)

    screenshot_box(doc, 'Captura: Tabla de calificaciones por actividad dentro de la dimensión SABER', height_cm=4.5)

    activity_box(doc,
        'Actividad 1.5 — Notas por actividad',
        [
            'Ingresar notas por actividad en al menos 2 dimensiones para 3 estudiantes.',
            'Verificar que la suma de actividades actualiza el total de la dimensión.',
            'Ver el historial de notas de un estudiante (historial por módulo y nivel).',
            'Exportar el historial de notas de un estudiante a PDF.',
        ]
    )

    page_break(doc)

    # 1.6
    heading_para(doc, '1.6  Gestión de Módulos y Contenido', level=1, color=TEAL, size=14, space_before=6)
    body_para(doc,
        'Como docente puedes crear y administrar el material educativo que tus estudiantes '
        'verán en sus módulos.')

    heading_para(doc, 'Estructura del contenido', level=2, color=BLUE, size=12, space_before=8)
    tip_box(doc,
        'Módulo → Lección → Sección. '
        'El módulo es la unidad temática del nivel, la lección agrupa secciones '
        'sobre un subtema, y la sección es el contenido mínimo: un texto, video, quiz, etc.',
        icon='\U0001f4d0', title='Estructura')

    heading_para(doc, 'Tipos de sección disponibles', level=2, color=BLUE, size=12, space_before=8)
    styled_table(doc,
        headers=['Tipo', 'Descripción'],
        rows=[
            ['Texto',          'Contenido escrito formateado'],
            ['Imagen',         'Imagen educativa con descripción'],
            ['Video',          'Video embebido desde URL (YouTube, etc.)'],
            ['Enlace',         'Enlace a recurso externo'],
            ['HTML',           'Contenido HTML personalizado'],
            ['Google Drive',   'Documento, PDF o presentación de Google Drive embebido'],
            ['Quiz',           'Preguntas de selección múltiple con puntaje e intentos'],
            ['Autoevaluación', 'Indicadores para que el estudiante se auto-evalúe (SER o DECIDIR)'],
        ],
        col_widths=[4.0, 11.5],
        font_size=10,
    )

    heading_para(doc, 'Asignar dimensión a una sección', level=2, color=BLUE, size=12, space_before=8)
    body_para(doc,
        'Cada sección debe estar vinculada a una dimensión evaluativa para que el progreso '
        'del estudiante contribuya a su nota en esa dimensión:')
    for dim in [
        'SABER — para lecturas, videos y materiales teóricos.',
        'HACER Proceso — para actividades prácticas de proceso.',
        'HACER Producto — para actividades orientadas al resultado final.',
        'Autoevaluación SER — para reflexiones sobre valores.',
        'Autoevaluación DECIDIR — para reflexiones sobre compromiso.',
    ]:
        bullet_para(doc, dim)

    heading_para(doc, 'Crear una lección con secciones', level=2, color=BLUE, size=12, space_before=8)
    for i, step in enumerate([
        'Seleccionar el módulo donde quieres agregar contenido.',
        'Hacer clic en “+ Nueva lección” y asignarle un título.',
        'Dentro de la lección, hacer clic en “+ Agregar sección”.',
        'Seleccionar el tipo de sección.',
        'Completar el contenido (texto, URL, preguntas, etc.).',
        'Asignar la dimensión evaluativa correspondiente.',
        'Guardar y activar la sección.',
    ], 1):
        numbered_step(doc, i, step)

    screenshot_box(doc, 'Captura: Gestor de contenido del módulo — lecciones y secciones', height_cm=4.5)

    activity_box(doc,
        'Actividad 1.6 — Crear contenido',
        [
            'Crear una lección nueva en tu módulo actual.',
            'Agregar una sección de tipo Texto con contenido sobre el tema.',
            'Crear un Quiz de 3 preguntas con 4 opciones cada una.',
            'Verificar que las secciones quedan activas y visibles.',
        ]
    )

    page_break(doc)

    # 1.7
    heading_para(doc, '1.7  Crear Material con IA y Publicarlo en la Plataforma',
                 level=1, color=TEAL, size=14, space_before=6)
    body_para(doc,
        'La Inteligencia Artificial puede ayudarte a preparar material educativo de calidad en menos tiempo. '
        'En esta sección aprenderás a combinar herramientas de IA con la plataforma del CEA.')

    heading_para(doc, 'Parte A — Crear material de estudio con IA', level=2, color=BLUE, size=12, space_before=8)
    body_para(doc, 'Herramientas recomendadas (versión gratuita es suficiente):', bold=True, size=10)
    styled_table(doc,
        headers=['Herramienta', 'Acceso', 'Descripción'],
        rows=[
            ['Claude',  'claude.ai',         'Muy preciso para textos educativos en español'],
            ['ChatGPT', 'chatgpt.com',        'Versátil, ideal para resúmenes y quizzes'],
            ['Gemini',  'gemini.google.com',  'Integrado con Google Drive y Docs'],
        ],
        col_widths=[3.0, 4.5, 8.0],
        font_size=10,
    )

    body_para(doc, 'Ejemplos de prompts efectivos:', bold=True, size=11, color=GREEN_D, space_after=4)

    for titulo, prompt_text in [
        ('Para una guía de estudio:',
         '“Crea una guía de estudio de 2 páginas sobre [tema] para estudiantes de nivel técnico básico. '
         'El lenguaje debe ser claro y sencillo. Incluye: introducción, conceptos principales con '
         'explicaciones breves, y un resumen final.”'),
        ('Para un resumen de proceso:',
         '“Explica paso a paso el proceso de [tema] de forma clara, numerada y con ejemplos simples '
         'para participantes que están aprendiendo por primera vez.”'),
        ('Para actividades prácticas:',
         '“Diseña 3 actividades prácticas sobre [tema] que puedan realizarse en un taller con materiales '
         'básicos. Incluye los materiales necesarios y el procedimiento.”'),
    ]:
        ptit = doc.add_paragraph()
        ptit.paragraph_format.space_before = Pt(4)
        ptit.paragraph_format.space_after  = Pt(2)
        ptit.paragraph_format.left_indent  = Cm(0.5)
        rtit = ptit.add_run(titulo + '  ')
        rtit.bold           = True
        rtit.font.name      = FONT_BODY
        rtit.font.size      = Pt(10)
        rtit.font.color.rgb = TEAL
        pprompt = doc.add_paragraph()
        pprompt.paragraph_format.space_before = Pt(0)
        pprompt.paragraph_format.space_after  = Pt(6)
        pprompt.paragraph_format.left_indent  = Cm(1.0)
        set_para_shading(pprompt, 'F0F4FA')
        rprompt = pprompt.add_run(prompt_text)
        rprompt.italic         = True
        rprompt.font.name      = FONT_BODY
        rprompt.font.size      = Pt(10)
        rprompt.font.color.rgb = GRAY_TEXT

    tip_box(doc,
        'Siempre revisa el material que genera la IA antes de publicarlo. '
        'Puede cometer errores en fechas o terminología local. Tú eres el experto en tu materia.',
        icon='⚠️', title='Importante')

    screenshot_box(doc, 'Captura: Interfaz de Claude o ChatGPT con un prompt de docente', height_cm=3.5)

    heading_para(doc, 'Parte B — Subir el material a Google Drive', level=2, color=BLUE, size=12, space_before=8)
    for i, step in enumerate([
        'Abrir Google Drive (drive.google.com) con tu cuenta de Gmail.',
        'Hacer clic en “+ Nuevo” → “Subir archivo” o crear un Documento de Google.',
        'Subir o escribir el documento con el material preparado.',
        'Hacer clic derecho sobre el archivo → “Compartir”.',
        'Cambiar el acceso a: “Cualquier persona con el enlace → Puede ver”.',
        'Hacer clic en “Copiar enlace” y guardar ese enlace para el siguiente paso.',
    ], 1):
        numbered_step(doc, i, step)

    screenshot_box(doc, 'Captura: Configuración de permisos de compartir en Google Drive', height_cm=3.5)

    heading_para(doc, 'Parte C — Publicar el material en la plataforma', level=2, color=BLUE, size=12, space_before=8)
    for i, step in enumerate([
        'En tu panel, ir a “Gestión de Contenido”.',
        'Seleccionar el módulo y la lección donde quieres agregar el material.',
        'Hacer clic en “+ Agregar sección” → seleccionar tipo “Google Drive”.',
        'Pegar el enlace copiado de Google Drive.',
        'El sistema extrae automáticamente el archivo y lo muestra integrado.',
        'Asignar la dimensión evaluativa (ej. SABER).',
        'Activar la sección y guardar.',
    ], 1):
        numbered_step(doc, i, step)
    tip_box(doc,
        'Los estudiantes verán el documento directamente dentro de la plataforma, '
        'sin necesidad de salir a otro sitio ni tener cuenta de Google.')

    heading_para(doc, 'Parte D — Crear un Quiz con IA', level=2, color=BLUE, size=12, space_before=8)
    body_para(doc, 'Prompt para generar el quiz:', bold=True, size=10)

    pquiz = doc.add_paragraph()
    pquiz.paragraph_format.left_indent  = Cm(0.5)
    pquiz.paragraph_format.space_after  = Pt(6)
    set_para_shading(pquiz, 'F0F4FA')
    rquiz = pquiz.add_run(
        '“Crea un quiz de 5 preguntas de selección múltiple sobre [tema]. '
        'Cada pregunta debe tener 4 opciones de respuesta (A, B, C, D). '
        'Indica cuál es la respuesta correcta. El nivel es técnico básico.”')
    rquiz.italic         = True
    rquiz.font.name      = FONT_BODY
    rquiz.font.size      = Pt(10)
    rquiz.font.color.rgb = GRAY_TEXT

    body_para(doc, 'Ingresar el quiz en la plataforma:', bold=True, size=10)
    for i, step in enumerate([
        'Crear una nueva sección de tipo Quiz en la lección.',
        'Hacer clic en “+ Agregar pregunta”.',
        'Escribir el enunciado y las opciones de respuesta.',
        'Marcar la opción correcta.',
        'Repetir para cada pregunta.',
        'Configurar el número de intentos permitidos (recomendado: 2 o 3).',
        'Guardar y activar.',
    ], 1):
        numbered_step(doc, i, step)

    screenshot_box(doc, 'Captura: Formulario de creación de Quiz en el gestor de contenido', height_cm=4.0)

    activity_box(doc,
        'Actividad 1.7 — Crear material con IA (práctica completa)',
        [
            'Elegir un tema de tu módulo actual.',
            'Con una herramienta de IA, generar una guía corta de 1 página sobre ese tema.',
            'Subir el documento a Google Drive y configurar permiso de visualización pública.',
            'Publicarlo en tu módulo como sección tipo Google Drive, dimensión SABER.',
            'Con la IA, generar un quiz de 5 preguntas sobre el mismo tema.',
            'Ingresar el quiz en la plataforma como sección tipo Quiz.',
            'Intercambiar módulo con un compañero y probar su quiz como estudiante.',
        ]
    )

    page_break(doc)

    # ══════════════════════════════════════════════════════════════════════════
    #  MODULO 2
    # ══════════════════════════════════════════════════════════════════════════
    section_divider(doc, 'MÓDULO 2 — Guía para Secretaria', BG_BLUE)

    body_para(doc,
        'La secretaria tiene acceso al panel administrativo en modo consulta. '
        'Puede ver toda la información y generar reportes PDF, pero no puede crear, editar ni eliminar usuarios.',
        italic=True, color=BLUE, space_after=8)

    # 2.1
    heading_para(doc, '2.1  El Panel de la Secretaria', level=1, color=BLUE, size=14, space_before=10)
    body_para(doc,
        'Al ingresar con tu cuenta de secretaria accederás al panel administrativo con las secciones: '
        'Carreras, Docentes, Estudiantes. Los botones de crear, editar y eliminar no estarán visibles.')

    screenshot_box(doc, 'Captura: Panel administrativo en vista de secretaria', height_cm=4.0)

    # 2.2
    heading_para(doc, '2.2  Consulta y Exportación de Estudiantes', level=1, color=BLUE, size=14, space_before=14)

    heading_para(doc, 'Buscar un estudiante', level=2, color=NAVY, size=12, space_before=8)
    body_para(doc, 'Ir a la sección “Estudiantes”. Puedes buscar por:', size=11)
    for item in ['Nombre o apellido', 'Código (ej. SIS-042)',
                 'N° de Carnet de Identidad', 'N° RUDEAL']:
        bullet_para(doc, item)
    body_para(doc,
        'También puedes filtrar por carrera y ordenar la lista haciendo clic en los encabezados de columna.')

    heading_para(doc, 'Ver el historial de notas', level=2, color=NAVY, size=12, space_before=8)
    for i, step in enumerate([
        'Hacer clic en el nombre del estudiante para abrir su perfil.',
        'Seleccionar “Ver notas” o el ícono de calificaciones.',
        'Se mostrará el historial por nivel y módulo con todas las dimensiones y el total.',
        'Hacer clic en “Exportar PDF” para generar el reporte imprimible.',
    ], 1):
        numbered_step(doc, i, step)

    heading_para(doc, 'Exportar lista de estudiantes a PDF', level=2, color=NAVY, size=12, space_before=8)
    for i, step in enumerate([
        'En la sección Estudiantes, hacer clic en el ícono de PDF junto a la carrera deseada.',
        'Seleccionar el turno (Tarde o Noche).',
        'Hacer clic en “Generar PDF”.',
        'El documento incluirá: encabezado institucional, carrera, turno, docentes asignados, '
        'total de participantes, lista agrupada por nivel (orden alfabético) y fecha.',
    ], 1):
        numbered_step(doc, i, step)

    screenshot_box(doc, 'Captura: Exportación de lista de estudiantes — selección de turno y carrera', height_cm=3.5)

    activity_box(doc,
        'Actividad 2.2 — Consulta y reportes',
        [
            'Buscar un estudiante de tu carrera por nombre y por código.',
            'Revisar su historial de notas completo.',
            'Exportar el historial de notas a PDF.',
            'Generar la lista de estudiantes de una carrera en turno tarde.',
        ]
    )

    # 2.3
    heading_para(doc, '2.3  Consulta y Exportación de Docentes', level=1, color=BLUE, size=14, space_before=14)
    body_para(doc,
        'En la sección “Docentes” verás la lista completa con: nombre, grado académico, carrera, turno y código. '
        'Usa la barra de búsqueda para filtrar. Para exportar, hacer clic en “Exportar PDF de docentes”. '
        'El documento incluirá apellidos, nombres, grado académico, carrera y turno.')

    # 2.4
    heading_para(doc, '2.4  Consulta de Carreras', level=1, color=BLUE, size=14, space_before=14)
    body_para(doc,
        'En la sección “Carreras” puedes ver todas las carreras activas, sus niveles y la cantidad '
        'de estudiantes matriculados por carrera. Útil para tener una visión general de la institución.')

    page_break(doc)

    # ══════════════════════════════════════════════════════════════════════════
    #  MODULO 3
    # ══════════════════════════════════════════════════════════════════════════
    section_divider(doc, 'MÓDULO 3 — Guía para Directora', BG_NAVY)

    body_para(doc,
        'La directora tiene acceso completo a todas las funciones del sistema.',
        italic=True, color=RGBColor(0x60, 0x80, 0xA0), space_after=8)

    # 3.1
    heading_para(doc, '3.1  Gestión de Carreras', level=1, color=NAVY, size=14, space_before=10)
    for i, step in enumerate([
        'Ir a la sección “Carreras” → “+ Nueva carrera”.',
        'Ingresar el nombre de la carrera (ej. Gastronomía).',
        'Ingresar el prefijo de código: entre 2 y 5 letras mayúsculas (ej. GAS). '
        'Se usará para los códigos de estudiantes: GAS-001, GAS-002…',
        'Guardar.',
    ], 1):
        numbered_step(doc, i, step)
    tip_box(doc,
        'Una carrera solo puede eliminarse si no tiene estudiantes registrados. '
        'Si tiene estudiantes, el sistema mostrará un mensaje de error.')

    # 3.2
    heading_para(doc, '3.2  Gestión de Docentes', level=1, color=NAVY, size=14, space_before=14)
    heading_para(doc, 'Crear un nuevo docente', level=2, color=BLUE, size=12, space_before=8)
    for i, step in enumerate([
        'Ir a “Docentes” → “+ Nuevo docente”.',
        'Completar: nombre(s), apellidos, celular (obligatorio), correo (opcional), '
        'grado académico, carrera y turno.',
        'La contraseña temporal se genera automáticamente.',
        'Hacer clic en “Crear docente”.',
        'Anotar el código generado (ej. DOC-005) y la contraseña para entregárselos al docente.',
    ], 1):
        numbered_step(doc, i, step)

    heading_para(doc, 'Resetear contraseña de un docente', level=2, color=BLUE, size=12, space_before=8)
    body_para(doc,
        'En la lista de docentes, buscar el ícono de “Resetear contraseña”, ingresar la nueva contraseña '
        'temporal, confirmar y entregar al docente para que la cambie en su siguiente ingreso.')

    screenshot_box(doc, 'Captura: Formulario de creación de docente con grado académico', height_cm=4.0)

    # 3.3
    heading_para(doc, '3.3  Gestión de Estudiantes', level=1, color=NAVY, size=14, space_before=14)
    heading_para(doc, 'Crear un nuevo estudiante', level=2, color=BLUE, size=12, space_before=8)
    body_para(doc, 'Ir a “Estudiantes” → “+ Nuevo estudiante”. Campos obligatorios:')
    for campo in [
        'Nombre(s) y apellidos (al menos uno).',
        'N° de Carnet de Identidad.',
        'Género (M / F).',
        'Fecha de nacimiento.',
        'Carrera, turno y nivel asignado.',
    ]:
        bullet_para(doc, campo)
    tip_box(doc,
        'El sistema valida la edad: bloquea el registro de menores de 14 años '
        'y muestra advertencia para participantes de 14 años.')

    heading_para(doc, 'Cambiar de nivel a un estudiante', level=2, color=BLUE, size=12, space_before=8)
    body_para(doc,
        'Abrir el perfil del estudiante → editar → cambiar el nivel → guardar. '
        'Usar cuando el participante avanza al siguiente nivel del programa.')

    # 3.4
    heading_para(doc, '3.4  Gestión de Personal Administrativo', level=1, color=NAVY, size=14, space_before=14)
    body_para(doc,
        'El sistema permite registrar una directora y una secretaria. '
        'Solo puede existir una cuenta activa de cada tipo simultáneamente.')
    for i, step in enumerate([
        'Ir a la sección “Administrativos”.',
        'Hacer clic en “+ Nueva secretaria” o “+ Nueva directora”.',
        'Completar: nombre(s), apellidos, celular (obligatorio), correo y grado académico (opcionales).',
        'La contraseña temporal se genera automáticamente.',
        'Anotar el código y contraseña para entregárselos.',
    ], 1):
        numbered_step(doc, i, step)

    # 3.5
    heading_para(doc, '3.5  Cuentas Bloqueadas', level=1, color=NAVY, size=14, space_before=14)
    body_para(doc,
        'El sistema bloquea automáticamente una cuenta tras varios intentos fallidos de ingreso. '
        'En la sección “Cuentas bloqueadas” verás: código, nombre, rol, fecha de bloqueo y número '
        'de intentos fallidos.')
    body_para(doc, 'Para desbloquear:', bold=True, size=11)
    for i, step in enumerate([
        'Localizar la cuenta en la lista.',
        'Hacer clic en “Desbloquear” y confirmar.',
        'Comunicar al usuario que su cuenta está activa y recomendarle cambiar su contraseña.',
    ], 1):
        numbered_step(doc, i, step)

    # 3.6
    heading_para(doc, '3.6  Configuración del Sistema', level=1, color=NAVY, size=14, space_before=14)
    styled_table(doc,
        headers=['Pestaña', 'Qué puedes configurar'],
        rows=[
            ['Institución', 'Nombre, misión, visión, teléfonos, correo y dirección'],
            ['Anuncio',     'Texto del anuncio y activar/desactivar en la página pública'],
            ['Semestre',    'Semestre activo visible en todo el sistema (ej. 1/2026)'],
            ['Galería',     'Imágenes de la página pública: agregar, reordenar (arrastre) y eliminar'],
            ['Requisitos',  'Requisitos de inscripción visibles para nuevos postulantes'],
        ],
        col_widths=[3.5, 12.0],
        font_size=10,
    )

    screenshot_box(doc, 'Captura: Panel de configuración del sistema — pestaña Institución', height_cm=4.0)

    activity_box(doc,
        'Actividades del Módulo 3',
        [
            'Registrar un docente nuevo con grado académico y exportar la lista PDF de docentes.',
            'Registrar un estudiante con todos los campos obligatorios.',
            'Crear la cuenta de secretaria para el taller con contraseña temporal.',
            'Activar un anuncio de prueba en la página pública y verificarlo.',
            'Actualizar el semestre activo al período vigente.',
            'Desbloquear una cuenta de prueba desde el panel.',
        ]
    )

    page_break(doc)

    # ══════════════════════════════════════════════════════════════════════════
    #  APENDICES
    # ══════════════════════════════════════════════════════════════════════════
    section_divider(doc, 'Apéndices', '374151')

    # A
    heading_para(doc, 'A.  Glosario', level=1, color=NAVY, size=14, space_before=10)
    styled_table(doc,
        headers=['Término', 'Definición'],
        rows=[
            ['Nivel',           'Etapa del programa (ej. Nivel 1, Nivel 2). Cada nivel contiene varios módulos'],
            ['Módulo',         'Unidad temática dentro de un nivel (ej. Módulo 3: Manipulación de Alimentos)'],
            ['Lección',         'Subdivisión de un módulo que agrupa secciones relacionadas'],
            ['Sección',         'Unidad mínima de contenido: un texto, video, quiz, etc.'],
            ['Dimensión',       'Área de evaluación: SER, SABER, HACER Proceso, HACER Producto, DECIDIR'],
            ['RUDEAL',          'Número de registro en el sistema RUDEAL del Ministerio de Educación'],
            ['Mesa Directiva',  'Estudiantes representantes. Tienen acceso ampliado a la asistencia'],
            ['Turno',           'Horario de clases: Tarde o Noche'],
            ['Grado académico', 'Título profesional: T.S., Lic., Ing., M.Sc., Dr.'],
            ['Contraseña temporal', 'Contraseña asignada al crear la cuenta. Debe cambiarse al primer ingreso'],
            ['Código de usuario',   'Identificador único (ej. SIS-042, DOC-005, ADM-DIR)'],
            ['Prompt',          'Instrucción que se da a una herramienta de Inteligencia Artificial'],
        ],
        col_widths=[3.8, 11.7],
        font_size=10,
    )

    # B
    heading_para(doc, 'B.  Tabla de Permisos por Rol', level=1, color=NAVY, size=14, space_before=14)
    styled_table(doc,
        headers=['Función', 'Docente', 'Secretaria', 'Directora'],
        rows=[
            ['Ver lista de estudiantes',     '✅ (su carrera)', '✅ (todos)',  '✅'],
            ['Registrar estudiantes',         '✅ (su carrera)', '❌',          '✅'],
            ['Editar / eliminar estudiantes', '❌',              '❌',          '✅'],
            ['Registrar asistencia',          '✅',              '❌',          '✅'],
            ['Ingresar calificaciones',       '✅',              '❌',          '✅'],
            ['Ver historial de notas',        '✅',              '✅',          '✅'],
            ['Exportar PDFs',                 '✅',              '✅',          '✅'],
            ['Gestionar contenido módulo',    '✅',              '❌',          '✅'],
            ['Ver / crear docentes',          '❌',              'Ver solo',    '✅'],
            ['Gestionar carreras',            '❌',              'Ver solo',    '✅'],
            ['Gestionar administrativos',     '❌',              '❌',          '✅'],
            ['Configuración del sistema',     '❌',              '❌',          '✅'],
            ['Desbloquear cuentas',           '❌',              '❌',          '✅'],
        ],
        col_widths=[5.5, 3.0, 3.0, 3.0],
        font_size=10,
    )

    # C
    heading_para(doc, 'C.  Rangos de Calificación por Dimensión', level=1, color=NAVY, size=14, space_before=14)
    styled_table(doc,
        headers=['Dimensión', 'Mínimo', 'Máximo', 'Observación'],
        rows=[
            ['SER',            '2',   '10',  'Calculado automáticamente desde asistencia'],
            ['SABER',          '6',   '30',  'Ingresado manualmente o por actividades'],
            ['HACER Proceso',  '6',   '30',  'Sugerido según progreso del estudiante en el módulo'],
            ['HACER Producto', '6',   '30',  'Ingresado manualmente o por actividades'],
            ['DECIDIR',        '2',   '10',  'Calculado automáticamente desde asistencia'],
            ['TOTAL',          '60',  '100', 'Suma de todas las dimensiones'],
        ],
        col_widths=[4.0, 2.0, 2.0, 7.5],
        font_size=10,
    )

    # D
    heading_para(doc, 'D.  Preguntas Frecuentes', level=1, color=NAVY, size=14, space_before=14)
    faqs = [
        ('¿Qué hago si olvidé mi contraseña?',
         'Comunicarte con tu docente (si eres estudiante) o con la directora '
         '(si eres docente o secretaria) para que reseteen tu contraseña.'),
        ('¿Por qué está bloqueada mi cuenta?',
         'El sistema bloquea una cuenta automáticamente tras varios intentos fallidos. '
         'La directora puede desbloquearla desde el panel administrativo.'),
        ('¿El sistema funciona desde el celular?',
         'Sí. La plataforma funciona en dispositivos móviles. Se recomienda Chrome o Firefox.'),
        ('¿Los cambios en asistencia o notas se guardan solos?',
         'La asistencia se guarda automáticamente al hacer clic. '
         'Las calificaciones se guardan al presionar el botón “Guardar”.'),
        ('¿Qué pasa si ingreso una nota fuera del rango?',
         'El campo rechaza el valor con una animación de sacudida y no guarda hasta '
         'que ingreses un número dentro del rango correcto.'),
        ('¿Los estudiantes pueden ver sus propias notas?',
         'Sí. Cada estudiante tiene acceso a su historial de notas desde su panel personal.'),
        ('¿Se puede compartir material de Drive sin que los estudiantes tengan Gmail?',
         'Sí. El material compartido como “cualquier persona con el enlace puede ver” '
         'es accesible para cualquier persona, sin cuenta de Google.'),
    ]
    for faq_q, faq_a in faqs:
        pq = doc.add_paragraph()
        pq.paragraph_format.space_before = Pt(6)
        pq.paragraph_format.space_after  = Pt(2)
        rq = pq.add_run('❓  ' + faq_q)
        rq.bold           = True
        rq.font.name      = FONT_BODY
        rq.font.size      = Pt(11)
        rq.font.color.rgb = NAVY
        pa = doc.add_paragraph()
        pa.paragraph_format.left_indent = Cm(0.6)
        pa.paragraph_format.space_after = Pt(2)
        ra = pa.add_run(faq_a)
        ra.font.name = FONT_BODY
        ra.font.size = Pt(11)

    # E
    page_break(doc)
    heading_para(doc, 'E.  Plantilla de Datos para Registro', level=1, color=NAVY, size=14, space_before=6)
    body_para(doc,
        'Usa estas plantillas para recopilar los datos en papel antes de registrarlos en el sistema.')

    heading_para(doc, 'Ficha de registro — Estudiante', level=2, color=BLUE, size=12, space_before=8)
    styled_table(doc,
        headers=['Campo', 'Dato'],
        rows=[
            ['Nombre(s)', ''],
            ['Apellido paterno', ''],
            ['Apellido materno', ''],
            ['N° Carnet de Identidad', ''],
            ['Fecha de nacimiento', 'dd/mm/aaaa'],
            ['Género', 'M  /  F'],
            ['N° RUDEAL (opcional)', ''],
            ['Celular', ''],
            ['Correo electrónico (opcional)', ''],
            ['Carrera', ''],
            ['Turno', 'Tarde  /  Noche'],
            ['Nivel', ''],
        ],
        col_widths=[6.5, 9.0],
        font_size=10,
    )

    doc.add_paragraph().paragraph_format.space_after = Pt(10)
    heading_para(doc, 'Ficha de registro — Docente', level=2, color=BLUE, size=12, space_before=8)
    styled_table(doc,
        headers=['Campo', 'Dato'],
        rows=[
            ['Nombre(s)', ''],
            ['Apellido paterno', ''],
            ['Apellido materno', ''],
            ['Grado académico', 'T.S.  /  Lic.  /  Ing.  /  M.Sc.  /  Dr.'],
            ['Celular', ''],
            ['Correo electrónico (opcional)', ''],
            ['Carrera', ''],
            ['Turno', 'Tarde  /  Noche'],
        ],
        col_widths=[6.5, 9.0],
        font_size=10,
    )

    doc.add_paragraph().paragraph_format.space_after = Pt(16)
    pf = doc.add_paragraph()
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rf = pf.add_run(
        'Guía elaborada para el taller de capacitación del CEA “Madre María Oliva” — Mayo 2026  ·  '
        'ceamadremariaoliva.edu.bo')
    rf.italic         = True
    rf.font.name      = FONT_BODY
    rf.font.size      = Pt(9)
    rf.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)

    return doc


if __name__ == '__main__':
    print('Generando documento Word...')
    doc = build_document()
    doc.save(OUT_PATH)
    print(f'Documento guardado en:\n   {OUT_PATH}')
